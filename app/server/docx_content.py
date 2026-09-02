from __future__ import annotations

from io import BytesIO
from pathlib import PurePosixPath
from zipfile import BadZipFile, ZipFile, is_zipfile
from xml.etree import ElementTree as ET

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
MAX_DOCX_ENTRIES = 500
MAX_DOCX_UNCOMPRESSED_BYTES = 20 * 1024 * 1024


class DocxContentError(ValueError):
    """Raised when a DOCX cannot be safely imported as training content."""


_FORBIDDEN_PACKAGE_PREFIXES = (
    "word/media/",
    "word/embeddings/",
    "word/activeX/",
)

_FORBIDDEN_XML_MARKERS = {
    b"<w:drawing": "Bilder, Screenshots oder Zeichnungen",
    b"<w:pict": "Bilder, Screenshots oder VML-Grafiken",
    b"<v:shape": "Formen oder Textfelder",
    b"<a:blip": "Bilder oder Grafiken",
    b"<c:chart": "Diagramme",
    b"<dgm:relIds": "SmartArt-Grafiken",
    b"<w:object": "eingebettete Objekte",
    b"<w:tbl": "Tabellen",
    b"<w:ins": "nachverfolgte Aenderungen",
    b"<w:del": "nachverfolgte Aenderungen",
    b"<w:hyperlink": "Hyperlinks",
    b"<w:footnoteReference": "Fussnoten",
    b"<w:endnoteReference": "Endnoten",
}


def validate_docx_package(data: bytes) -> None:
    if not data or not is_zipfile(BytesIO(data)):
        raise DocxContentError("Die Datei ist kein gueltiges DOCX-Dokument.")

    try:
        with ZipFile(BytesIO(data)) as archive:
            infos = archive.infolist()
            if len(infos) > MAX_DOCX_ENTRIES:
                raise DocxContentError("Das DOCX enthaelt zu viele interne Dateien.")

            total_uncompressed = 0
            for info in infos:
                normalized = info.filename.replace("\\", "/")
                path = PurePosixPath(normalized)
                if path.is_absolute() or ".." in path.parts:
                    raise DocxContentError("Das DOCX enthaelt einen unzulaessigen Dateipfad.")
                total_uncompressed += info.file_size
                if total_uncompressed > MAX_DOCX_UNCOMPRESSED_BYTES:
                    raise DocxContentError("Das DOCX ist intern zu gross und wird aus Sicherheitsgruenden abgelehnt.")
                if any(normalized.startswith(prefix) for prefix in _FORBIDDEN_PACKAGE_PREFIXES):
                    raise DocxContentError(
                        "DOCX-Import abgelehnt: Bilder, Screenshots, eingebettete Dateien oder aktive Inhalte sind nicht zulaessig."
                    )

            document_xml = archive.read("word/document.xml")
            for marker, description in _FORBIDDEN_XML_MARKERS.items():
                if marker in document_xml:
                    if marker in {b"<w:drawing", b"<w:pict", b"<v:shape", b"<a:blip"}:
                        raise DocxContentError(
                            "DOCX-Import abgelehnt: Bilder, Screenshots, Grafiken, Formen und Textfelder sind nicht zulaessig."
                        )
                    raise DocxContentError(f"DOCX-Import abgelehnt: {description} werden nicht unterstuetzt.")
    except KeyError as error:
        raise DocxContentError("Die DOCX-Datei enthaelt kein gueltiges Word-Dokument.") from error
    except BadZipFile as error:
        raise DocxContentError("Die Datei ist kein gueltiges DOCX-Dokument.") from error


def import_training_content_docx(data: bytes) -> str:
    validate_docx_package(data)
    try:
        document = Document(BytesIO(data))
    except Exception as error:  # python-docx exposes multiple parser exception types
        raise DocxContentError("Das DOCX konnte nicht gelesen werden.") from error

    numbering = _numbering_formats(data)
    own_export = (document.core_properties.comments or "").startswith("Exported by Schulungsplantool")
    lines: list[str] = []
    for index, paragraph in enumerate(document.paragraphs):
        if own_export and index == 0 and paragraph.text.startswith("Schulungspunkte - "):
            continue
        text = _paragraph_to_markdown(paragraph)
        if not text:
            if lines and lines[-1] != "":
                lines.append("")
            continue

        heading_level = _heading_level(paragraph)
        if heading_level:
            lines.append(f"{'#' * heading_level} {text}")
            continue

        list_kind = _list_kind(paragraph, numbering)
        if list_kind == "bullet":
            lines.append(f"- {text}")
        elif list_kind == "number":
            lines.append(f"1. {text}")
        else:
            lines.append(text)

    return _collapse_blank_lines("\n".join(lines).strip())


def export_training_content_docx(title: str, markdown_content: str) -> bytes:
    document = Document()
    document.core_properties.title = f"Schulungspunkte - {title}"
    document.core_properties.author = "Schulungsplantool"
    document.core_properties.comments = "Exported by Schulungsplantool for local DOCX round-trip editing."

    _add_usage_notice_to_header(document)

    title_paragraph = document.add_heading(f"Schulungspunkte - {title}", level=1)
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    for raw_line in str(markdown_content or "").replace("\r\n", "\n").replace("\r", "\n").split("\n"):
        line = raw_line.rstrip()
        if not line.strip():
            document.add_paragraph("")
            continue

        heading = _parse_heading_line(line)
        if heading is not None:
            level, content = heading
            paragraph = document.add_heading(level=level)
            _add_markdown_inline_runs(paragraph, content)
            continue

        bullet = _parse_bullet_line(line)
        if bullet is not None:
            paragraph = document.add_paragraph(style="List Bullet")
            _add_markdown_inline_runs(paragraph, bullet)
            continue

        numbered = _parse_numbered_line(line)
        if numbered is not None:
            paragraph = document.add_paragraph(style="List Number")
            _add_markdown_inline_runs(paragraph, numbered)
            continue

        paragraph = document.add_paragraph()
        _add_markdown_inline_runs(paragraph, line)

    stream = BytesIO()
    document.save(stream)
    return stream.getvalue()


def _add_usage_notice_to_header(document: Document) -> None:
    section = document.sections[0]
    header = section.header
    heading = header.paragraphs[0]
    heading.text = "Hinweis zur Bearbeitung und zum Re-Import"
    heading.runs[0].bold = True
    heading.runs[0].font.size = Pt(9)

    allowed = header.add_paragraph()
    run = allowed.add_run(
        "Zulässig: Überschriften Ebene 1-3, normaler Text, Fett, Kursiv, Aufzählungen und Nummerierungen."
    )
    run.font.size = Pt(8)

    forbidden = header.add_paragraph()
    run = forbidden.add_run(
        "Nicht zulässig: Bilder oder Screenshots, Grafiken/Formen/Textfelder, Tabellen, Diagramme/SmartArt, "
        "eingebettete Dateien, Hyperlinks, Fuß-/Endnoten und nicht angenommene nachverfolgte Änderungen. "
        "Dokumente mit Bildern oder Screenshots werden beim Import vollständig abgelehnt."
    )
    run.bold = True
    run.font.size = Pt(8)


def _collapse_blank_lines(value: str) -> str:
    if not value:
        return value
    output: list[str] = []
    newline_run = 0
    for char in value:
        if char == "\n":
            newline_run += 1
            if newline_run <= 2:
                output.append(char)
        else:
            newline_run = 0
            output.append(char)
    return "".join(output)


def _skip_whitespace(value: str, position: int) -> int:
    while position < len(value) and value[position].isspace():
        position += 1
    return position


def _parse_heading_line(line: str) -> tuple[int, str] | None:
    level = 0
    while level < len(line) and level < 4 and line[level] == "#":
        level += 1
    if level not in {1, 2, 3} or level >= len(line) or not line[level].isspace():
        return None
    content_start = _skip_whitespace(line, level)
    if content_start >= len(line):
        return None
    return level, line[content_start:]


def _parse_bullet_line(line: str) -> str | None:
    if len(line) < 2 or line[0] not in {"-", "*"} or not line[1].isspace():
        return None
    content_start = _skip_whitespace(line, 1)
    if content_start >= len(line):
        return None
    return line[content_start:]


def _parse_numbered_line(line: str) -> str | None:
    position = 0
    while position < len(line) and line[position].isdigit():
        position += 1
    if position == 0 or position >= len(line) or line[position] != ".":
        return None
    position += 1
    if position >= len(line) or not line[position].isspace():
        return None
    content_start = _skip_whitespace(line, position)
    if content_start >= len(line):
        return None
    return line[content_start:]


def _next_inline_marker(value: str, position: int) -> int:
    star = value.find("*", position)
    code = value.find("`", position)
    if star == -1:
        return code
    if code == -1:
        return star
    return min(star, code)


def _add_markdown_inline_runs(paragraph, value: str) -> None:
    position = 0
    while position < len(value):
        marker_start = _next_inline_marker(value, position)
        if marker_start == -1:
            paragraph.add_run(value[position:])
            return
        if marker_start > position:
            paragraph.add_run(value[position:marker_start])

        if value[marker_start] == "`":
            delimiter = "`"
        elif value.startswith("***", marker_start):
            delimiter = "***"
        elif value.startswith("**", marker_start):
            delimiter = "**"
        else:
            delimiter = "*"

        content_start = marker_start + len(delimiter)
        marker_end = value.find(delimiter, content_start)
        if marker_end <= content_start:
            paragraph.add_run(value[marker_start:])
            return

        run = paragraph.add_run(value[content_start:marker_end])
        if delimiter == "***":
            run.bold = True
            run.italic = True
        elif delimiter == "**":
            run.bold = True
        elif delimiter == "*":
            run.italic = True
        else:
            run.font.name = "Consolas"
        position = marker_end + len(delimiter)


def _paragraph_to_markdown(paragraph) -> str:
    parts: list[str] = []
    for run in paragraph.runs:
        text = run.text.replace("\r", " ").replace("\n", " ")
        if not text:
            continue
        text = _escape_markdown_text(text)
        if run.bold and run.italic:
            parts.append(f"***{text}***")
        elif run.bold:
            parts.append(f"**{text}**")
        elif run.italic:
            parts.append(f"*{text}*")
        else:
            parts.append(text)
    return "".join(parts).strip()


def _escape_markdown_text(value: str) -> str:
    return value.replace("\\", "\\\\").replace("*", "\\*")


def _heading_level(paragraph) -> int | None:
    style = paragraph.style
    style_id = (getattr(style, "style_id", "") or "").lower()
    style_name = (getattr(style, "name", "") or "").lower()
    for level in (1, 2, 3):
        if style_id in {f"heading{level}", f"berschrift{level}"}:
            return level
        compact_name = style_name.rstrip()
        if compact_name.endswith(str(level)):
            prefix = compact_name[:-1].rstrip()
            if prefix in {"heading", "ueberschrift", "überschrift"}:
                return level
    return None


def _list_kind(paragraph, numbering: dict[tuple[int, int], str]) -> str | None:
    p_pr = paragraph._p.pPr
    if p_pr is not None and p_pr.numPr is not None and p_pr.numPr.numId is not None:
        num_id = int(p_pr.numPr.numId.val)
        ilvl = int(p_pr.numPr.ilvl.val) if p_pr.numPr.ilvl is not None else 0
        fmt = numbering.get((num_id, ilvl)) or numbering.get((num_id, 0), "")
        return "bullet" if fmt == "bullet" else "number"

    style_id = (getattr(paragraph.style, "style_id", "") or "").lower()
    style_name = (getattr(paragraph.style, "name", "") or "").lower()
    if "listbullet" in style_id or "aufzaehlung" in style_name or "aufzählung" in style_name:
        return "bullet"
    if "listnumber" in style_id or "nummer" in style_name:
        return "number"
    return None


def _numbering_formats(data: bytes) -> dict[tuple[int, int], str]:
    try:
        with ZipFile(BytesIO(data)) as archive:
            xml = archive.read("word/numbering.xml")
    except (KeyError, BadZipFile):
        return {}

    namespace = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    value_attr = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val"
    root = ET.fromstring(xml)
    abstract_formats: dict[tuple[int, int], str] = {}
    for abstract in root.findall("w:abstractNum", namespace):
        abstract_id = int(abstract.attrib.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}abstractNumId", "0"))
        for level in abstract.findall("w:lvl", namespace):
            ilvl = int(level.attrib.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}ilvl", "0"))
            fmt = level.find("w:numFmt", namespace)
            if fmt is not None:
                abstract_formats[(abstract_id, ilvl)] = fmt.attrib.get(value_attr, "")

    result: dict[tuple[int, int], str] = {}
    for num in root.findall("w:num", namespace):
        num_id = int(num.attrib.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numId", "0"))
        abstract_ref = num.find("w:abstractNumId", namespace)
        if abstract_ref is None:
            continue
        abstract_id = int(abstract_ref.attrib.get(value_attr, "0"))
        for (candidate_abstract_id, ilvl), fmt in abstract_formats.items():
            if candidate_abstract_id == abstract_id:
                result[(num_id, ilvl)] = fmt
    return result
