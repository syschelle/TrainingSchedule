from base64 import b64decode
from io import BytesIO
from zipfile import ZipFile

import pytest
from docx import Document

from app.server.docx_content import (
    DocxContentError,
    export_training_content_docx,
    import_training_content_docx,
    validate_docx_package,
)


SINGLE_PIXEL_PNG = b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAusB9Y9ZP8sAAAAASUVORK5CYII="
)


def _save(document: Document) -> bytes:
    stream = BytesIO()
    document.save(stream)
    return stream.getvalue()


def test_docx_export_contains_usage_notice_and_roundtrips_markdown():
    markdown = "## Grundlagen\n\n- Anmeldung\n- **Patientensuche**\n\n1. Viewer oeffnen\n2. *Messwerkzeuge*"
    data = export_training_content_docx("DU Diagnost Basic", markdown)

    validate_docx_package(data)
    with ZipFile(BytesIO(data)) as archive:
        assert not any(name.startswith("word/media/") for name in archive.namelist())

    document = Document(BytesIO(data))
    header_text = "\n".join(paragraph.text for paragraph in document.sections[0].header.paragraphs)
    assert "Hinweis zur Bearbeitung und zum Re-Import" in header_text
    assert "Bilder oder Screenshots" in header_text
    assert "vollständig abgelehnt" in header_text

    imported = import_training_content_docx(data)
    assert "# Schulungspunkte - DU Diagnost Basic" not in imported
    assert "## Grundlagen" in imported
    assert "- Anmeldung" in imported
    assert "- **Patientensuche**" in imported
    assert "1. Viewer oeffnen" in imported
    assert "1. *Messwerkzeuge*" in imported


def test_docx_import_rejects_any_embedded_image_or_screenshot(tmp_path):
    document = Document()
    document.add_paragraph("Schulungspunkt vor dem Screenshot")
    image_path = tmp_path / "screenshot.png"
    image_path.write_bytes(SINGLE_PIXEL_PNG)
    document.add_picture(str(image_path))

    with pytest.raises(DocxContentError, match="Bilder, Screenshots"):
        import_training_content_docx(_save(document))


def test_docx_import_rejects_tables():
    document = Document()
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Punkt"
    table.cell(0, 1).text = "Beschreibung"

    with pytest.raises(DocxContentError, match="Tabellen"):
        import_training_content_docx(_save(document))


def test_docx_import_accepts_supported_word_content():
    document = Document()
    document.add_heading("Grundlagen", level=2)
    paragraph = document.add_paragraph()
    paragraph.add_run("Fett").bold = True
    paragraph.add_run(" und ")
    paragraph.add_run("kursiv").italic = True
    document.add_paragraph("Erster Punkt", style="List Bullet")
    document.add_paragraph("Erster Schritt", style="List Number")

    imported = import_training_content_docx(_save(document))
    assert "## Grundlagen" in imported
    assert "**Fett** und *kursiv*" in imported
    assert "- Erster Punkt" in imported
    assert "1. Erster Schritt" in imported
